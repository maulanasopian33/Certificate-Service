from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

def fill_docx(template_path, output_path, context, qr_path):
    doc = Document(template_path)

    def process_paragraph(p):
        if "{{qr_code}}" in p.text:
            p.text = ""
            run = p.add_run()
            run.add_picture(str(qr_path), width=Inches(1.5))
            return

        for key, val in context.items():
            if f"{{{{{key}}}}}" in p.text:
                p.text = p.text.replace(f"{{{{{key}}}}}", str(val))

    for p in doc.paragraphs:
        process_paragraph(p)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    process_paragraph(p)

    # Process Text Boxes (w:txbxContent)
    for txbx in doc.element.body.xpath('//w:txbxContent'):
        for p_element in txbx.findall(qn('w:p')):
            process_paragraph(Paragraph(p_element, doc))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
